"""Parser for DOCX files using python-docx.

Extracts paragraph text and tries to find a title from either core properties
or the first Heading 1 style paragraph.
"""

from __future__ import annotations

from pathlib import Path

from .base import (
    DocumentParser,
    ParsedDocument,
    ParserError,
    compute_hash,
    fallback_title_from_path,
    normalize_text,
    chunk_text_sliding,
)


class WordParser(DocumentParser):
    def can_parse(self, suffix: str) -> bool:
        return suffix.lower() == ".docx"

    def parse(self, path: Path, chunk_size: int = 1500, overlap: int = 200) -> ParsedDocument:
        try:
            from docx import Document  # python-docx
        except ImportError as e:
            raise ParserError(
                "python-docx is not installed. Run `uv pip install -e .`."
            ) from e

        try:
            doc = Document(str(path))
        except Exception as e:
            raise ParserError(f"Cannot open DOCX {path.name}: {e}") from e

        # Collect paragraph text, preserving heading style markers so the LLM
        # in Stage 3 still has structural signal.
        lines: list[str] = []
        title_from_heading: str | None = None

        for para in doc.paragraphs:
            text_line = para.text.strip()
            if not text_line:
                continue
            style_name = para.style.name if para.style else ""
            if style_name.startswith("Heading"):
                # Prepend hashes to make it markdown-ish
                try:
                    level = int(style_name.replace("Heading", "").strip() or "1")
                except ValueError:
                    level = 1
                level = max(1, min(level, 6))
                lines.append("#" * level + " " + text_line)
                if title_from_heading is None and level == 1:
                    title_from_heading = text_line
            else:
                lines.append(text_line)

        # Also extract tables (flatten to pipe-separated rows)
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    lines.append(" | ".join(cells))

        full_text = "\n".join(lines)
        text = normalize_text(full_text)

        # Title extraction: core properties > first Heading 1 > filename
        title: str | None = None
        try:
            core_title = doc.core_properties.title
            if core_title:
                title = core_title.strip()
        except Exception:
            pass
        if not title:
            title = title_from_heading
        if not title:
            title = fallback_title_from_path(path)

        # Metadata
        metadata: dict = {"paragraph_count": len(doc.paragraphs)}
        try:
            props = doc.core_properties
            if props.author:
                metadata["author"] = props.author
            if props.created:
                metadata["created"] = str(props.created)
        except Exception:
            pass

        chunks = chunk_text_sliding(text, chunk_size, overlap)
        return ParsedDocument(
            source_path=path,
            file_type="docx",
            title=title,
            text=text,
            content_hash=compute_hash(text),
            bytes=path.stat().st_size,
            metadata=metadata,
            chunks=chunks,
        )


def parse(path: Path) -> ParsedDocument:
    """Parse a .docx file."""
    return WordParser().parse(path)
